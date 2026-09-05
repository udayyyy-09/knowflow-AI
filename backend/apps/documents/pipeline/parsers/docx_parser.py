"""
DOCX Document Parser using python-docx.
Extracts paragraphs and headings, preserving document structural hierarchy.
"""
import io
import logging
from typing import Union, BinaryIO
from docx import Document as DocxDocument
from apps.documents.pipeline.parsers.base import BaseParser, ParsedDocument, ParsedBlock

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    Parser for Microsoft Word (.docx) documents.
    Extracts text while preserving section headings.
    """

    def parse(self, file_path_or_obj: Union[str, BinaryIO, bytes], original_filename: str = "") -> ParsedDocument:
        """
        Parses a .docx file and returns structured blocks grouped by section headers.
        """
        if isinstance(file_path_or_obj, bytes):
            stream = io.BytesIO(file_path_or_obj)
        elif isinstance(file_path_or_obj, str):
            stream = file_path_or_obj
        else:
            stream = file_path_or_obj
            if hasattr(stream, 'seek'):
                stream.seek(0)

        doc = DocxDocument(stream)
        blocks = []
        current_header = ""
        current_paragraphs = []
        total_chars = 0
        total_words = 0

        def flush_current_block():
            nonlocal current_paragraphs, total_chars, total_words
            if current_paragraphs:
                block_text = "\n\n".join(current_paragraphs).strip()
                if block_text:
                    c_count = len(block_text)
                    w_count = len(block_text.split())
                    total_chars += c_count
                    total_words += w_count
                    blocks.append(
                        ParsedBlock(
                            content=block_text,
                            page_number=None,
                            section_header=current_header or "Main",
                            metadata={
                                "section_header": current_header or "Main",
                                "char_count": c_count,
                                "word_count": w_count,
                            }
                        )
                    )
                current_paragraphs = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name.lower() if p.style and p.style.name else ""
            if "heading" in style_name or style_name.startswith("title"):
                flush_current_block()
                current_header = text
            else:
                current_paragraphs.append(text)

        flush_current_block()

        # Extract table text if any
        for t_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                table_text = "\n".join(table_rows)
                c_count = len(table_text)
                w_count = len(table_text.split())
                total_chars += c_count
                total_words += w_count
                blocks.append(
                    ParsedBlock(
                        content=table_text,
                        page_number=None,
                        section_header=f"Table {t_idx + 1}",
                        metadata={"is_table": True, "table_index": t_idx + 1}
                    )
                )

        # Core properties metadata
        meta = {}
        try:
            core = doc.core_properties
            if core.title:
                meta["title"] = core.title
            if core.author:
                meta["author"] = core.author
            if core.created:
                meta["created"] = str(core.created)
        except Exception:
            pass

        return ParsedDocument(
            blocks=blocks,
            total_pages=None,
            char_count=total_chars,
            word_count=total_words,
            title=meta.get("title") or None,
            metadata=meta
        )
